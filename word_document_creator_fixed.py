#!/usr/bin/env python3
"""
Fixed Word Document Creator for Financial Analysis Reports
Addresses: Remove PE columns, fix EPS units, remove duplicate data, format calculations in table
"""

import os
import json
from datetime import datetime
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.patches import Rectangle
import numpy as np
from io import BytesIO
from config_manager import ConfigManager
import re
from docx.shared import RGBColor

def clean_agent_message(message: str) -> str:
    """Clean agent message to remove thinking process and extract only meaningful content"""
    if not message:
        return "No message available"
    
    import re
    
    # Store original message for debugging
    original_message = message
    
    # Remove thinking process tags and content (more aggressive)
    message = re.sub(r'<think>.*?</think>', '', message, flags=re.DOTALL | re.IGNORECASE)
    message = re.sub(r'<thinking>.*?</thinking>', '', message, flags=re.DOTALL | re.IGNORECASE)
    message = re.sub(r'<thought>.*?</thought>', '', message, flags=re.DOTALL | re.IGNORECASE)
    message = re.sub(r'<reasoning>.*?</reasoning>', '', message, flags=re.DOTALL | re.IGNORECASE)
    message = re.sub(r'<analysis>.*?</analysis>', '', message, flags=re.DOTALL | re.IGNORECASE)
    message = re.sub(r'<reason>.*?</reason>', '', message, flags=re.DOTALL | re.IGNORECASE)
    message = re.sub(r'<process>.*?</process>', '', message, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove agent prefixes
    message = re.sub(r'^(Financial|Research|Strategic)\s+Analyst:\s*', '', message, flags=re.IGNORECASE)
    
    # Remove any remaining thinking indicators
    message = re.sub(r'Thought:\s*.*?(?=\n|$)', '', message, flags=re.DOTALL | re.IGNORECASE)
    message = re.sub(r'Thinking:\s*.*?(?=\n|$)', '', message, flags=re.DOTALL | re.IGNORECASE)
    message = re.sub(r'Analysis:\s*.*?(?=\n|$)', '', message, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove markdown code blocks
    message = re.sub(r'```.*?```', '', message, flags=re.DOTALL)
    
    # Extract JSON if present
    json_match = re.search(r'\{[\s\S]*\}', message)
    if json_match:
        try:
            json_data = json.loads(json_match.group(0))
            # Convert JSON to readable format
            return json.dumps(json_data, indent=2)
        except:
            pass
    
    # Clean up extra whitespace and newlines
    message = re.sub(r'\n+', '\n', message)
    message = re.sub(r'\s+', ' ', message)
    message = message.strip()
    
    # If message is too long, truncate it
    if len(message) > 1000:
        message = message[:1000] + "..."
    
    # If message is empty after cleaning, return a default message
    if not message or message.strip() == "":
        return "Agent output cleaned - no meaningful content found"
    
    return message

def create_word_document(
    eps_data: Dict[str, Any],
    roe_data: Dict[str, Any],
    market_data: Dict[str, Any],
    calculations: Dict[str, Any],
    valuation_data: Dict[str, Any],
    research_data: Optional[Dict[str, Any]] = None,
    validation_results: Optional[Dict[str, Any]] = None,
    filename: Optional[str] = None
) -> str:
    """
    Create a comprehensive Word document with financial analysis
    
    Args:
        eps_data: EPS data dictionary
        roe_data: ROE data dictionary
        market_data: Market data dictionary
        calculations: Calculations dictionary
        valuation_data: Valuation data dictionary
        research_data: Research data dictionary (optional, will be generated dynamically if not provided)
        validation_results: Validation results and recommendations (optional)
        filename: Output filename (optional)
    
    Returns:
        Path to the generated Word document
    """
    
    # Initialize configuration manager
    config_manager = ConfigManager()
    company_info = config_manager.get_company_info()
    company_name = company_info.get('name', 'Company')
    ticker = company_info.get('ticker', 'TICKER')
    industry = company_info.get('industry', 'Industry')
    country = company_info.get('country', 'Country')
    
    # Initialize research_data if not provided
    if research_data is None:
        research_data = {}
    
    # Initialize validation_results if not provided
    if validation_results is None:
        validation_results = {}
    
    # Create document
    doc = Document()
    
    # Set up document styles
    setup_document_styles(doc)
    
    # Add title page
    add_title_page(doc, company_name, company_info)
    
    # Add executive summary
    add_executive_summary(doc, company_name, market_data, research_data, company_info)
    
    # Add market sentiment analysis
    add_market_sentiment_analysis(doc, company_name, industry, research_data, company_info)
    
    # Add SWOT analysis
    add_swot_analysis(doc, company_name, industry, research_data, company_info)
    
    # Add financial analysis
    add_financial_analysis(doc, eps_data, roe_data, calculations, company_name, company_info)
    
    # Add validation results and recommendations
    # Merge Data Validation Results & Recommendations with Valuation Analysis
    add_validation_and_valuation(doc, validation_results, valuation_data, calculations, market_data, company_name, company_info, eps_data)
    
    # Add market analysis
    add_market_analysis(doc, company_name, ticker, industry, research_data, company_info)
    
    # Add conclusion
    add_conclusion(doc, company_name, industry, research_data, company_info)
    
    # Add sources
    add_sources(doc, company_name, industry, research_data, company_info, eps_data)
    
    # Generate filename if not provided
    if filename is None:
        filename_prefix = config_manager.get_filename_prefix()
        filename = f"{filename_prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    # Save document
    doc.save(filename)
    print(f"Word document created: {filename}")
    
    return filename

def calculate_financial_metrics(eps_data, units):
    """
    Calculate financial metrics from EPS data
    """
    # Filter out non-year keys and get valid EPS data
    eps_values = []
    years = []
    
    for year in sorted(eps_data.keys()):
        if year.isdigit():
            year_data = eps_data[year]
            if isinstance(year_data, dict) and year_data.get('basic_eps') is not None:
                eps_values.append(year_data['basic_eps'])
                years.append(int(year))
    
    if len(eps_values) < 2:
        return {}
    
    # Calculate CAGR
    initial_eps = eps_values[0]
    final_eps = eps_values[-1]
    num_years = years[-1] - years[0]
    
    if initial_eps > 0 and num_years > 0:
        cagr = ((final_eps / initial_eps) ** (1 / num_years) - 1) * 100
    else:
        cagr = None
    
    # Get current EPS (most recent)
    current_eps = eps_values[-1] if eps_values else None
    
    # Calculate projected EPS for 10 years
    if current_eps is not None and cagr is not None:
        projected_eps_10yr = current_eps * (1 + cagr / 100) ** 10
    else:
        projected_eps_10yr = None
    
    # Calculate average and lowest PE ratios (no hardcoded fallback; leave None if unavailable)
    pe_values = []
    for year_data in eps_data.values():
        if isinstance(year_data, dict) and year_data.get('pe_ratio') is not None:
            pe_values.append(year_data['pe_ratio'])
    
    if pe_values:
        avg_pe_ratio = sum(pe_values) / len(pe_values)
        lowest_pe_ratio = min(pe_values)
    else:
        avg_pe_ratio = None
        lowest_pe_ratio = None
    
    return {
        'cagr': cagr,
        'avg_pe_ratio': avg_pe_ratio,
        'lowest_pe_ratio': lowest_pe_ratio,
        'current_eps': current_eps,
        'projected_eps_10yr': projected_eps_10yr
    }

def create_eps_graph(eps_data, units):
    """
    Create EPS trend graph
    """
    try:
        # Filter out non-year keys and get valid data
        years = []
        eps_values = []
        
        for year in sorted(eps_data.keys()):
            if year.isdigit():
                year_data = eps_data[year]
                if isinstance(year_data, dict) and year_data.get('basic_eps') is not None:
                    years.append(int(year))
                    eps_values.append(year_data['basic_eps'])
        
        if len(years) < 2:
            return None
        
        # Create the plot
        plt.figure(figsize=(10, 6))
        plt.plot(years, eps_values, marker='o', linewidth=2, markersize=6)
        
        # Customize the plot
        plt.title('Historical Basic Earnings Per Share Trend', fontsize=14, fontweight='bold')
        plt.xlabel('Year', fontsize=12)
        plt.ylabel('Basic EPS (SGD)', fontsize=12)
        plt.grid(True, alpha=0.3)
        
        # Format y-axis to show currency
        eps_unit = units.get('eps_unit', 'cents')
        if eps_unit == 'cents':
            plt.ylabel('Basic EPS (SGD cents)', fontsize=12)
        else:
            plt.ylabel('Basic EPS (SGD)', fontsize=12)
        
        # Rotate x-axis labels for better readability
        plt.xticks(years, rotation=45)
        
        # Add value labels on points
        for i, (year, value) in enumerate(zip(years, eps_values)):
            plt.annotate(f'{value:.2f}', (year, value), 
                        textcoords="offset points", xytext=(0,10), ha='center')
        
        plt.tight_layout()
        
        # Save to buffer
        buffer = BytesIO()
        plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
        buffer.seek(0)
        plt.close()
        
        return buffer
        
    except Exception as e:
        print(f"Error creating EPS graph: {e}")
        return None

def create_roe_graph(roe_data, units):
    """
    Create ROE trend graph
    """
    try:
        # Filter out non-year keys and get valid ROE data
        roe_values = []
        years = []
        
        for year in sorted(roe_data.keys()):
            if not year.isdigit():
                continue
            year_data = roe_data[year]
            if isinstance(year_data, dict) and year_data.get('basic_roe') is not None:
                roe_values.append(year_data['basic_roe'])
                years.append(int(year))
            elif isinstance(year_data, (int, float)) and year_data is not None:
                roe_values.append(year_data)
                years.append(int(year))
        
        if len(roe_values) < 2:
            print("Insufficient ROE data for graph creation")
            return None
        
        # Create the graph
        plt.figure(figsize=(10, 6))
        plt.plot(years, roe_values, marker='s', linewidth=2, markersize=6, color='green')
        plt.title('ROE Historical Trend', fontsize=14, fontweight='bold')
        plt.xlabel('Year', fontsize=12)
        plt.ylabel('ROE (%)', fontsize=12)
        plt.grid(True, alpha=0.3)
        plt.xticks(years, rotation=45)
        
        # Add value labels on points
        for i, (year, value) in enumerate(zip(years, roe_values)):
            plt.annotate(f'{value:.2f}%', (year, value), 
                        textcoords="offset points", xytext=(0,10), ha='center',
                        fontsize=9, fontweight='bold')
        
        plt.tight_layout()
        
        # Save to buffer
        buffer = BytesIO()
        plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
        buffer.seek(0)
        plt.close()
        
        return buffer
        
    except Exception as e:
        print(f"Error creating ROE graph: {e}")
        return None

def setup_document_styles(doc: Document):
    """Set up document styles for consistent formatting"""
    # Title style
    title_style = doc.styles['Title']
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    # Heading 1 style
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.name = 'Arial'
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True
    
    # Heading 2 style
    heading2_style = doc.styles['Heading 2']
    heading2_style.font.name = 'Arial'
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True
    
    # Normal style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)

def add_title_page(doc: Document, company_name: str, company_info: Dict[str, Any]):
    """Add title page to the document"""
    # Title
    title = doc.add_heading(f'{company_name} - Financial Analysis Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('Comprehensive Financial Performance Analysis').italic = True
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
    
    doc.add_page_break()

def add_executive_summary(doc: Document, company_name: str, market_data: Dict[str, Any], research_data: Dict[str, Any], company_info: Dict[str, Any]):
    """Add executive summary section"""
    doc.add_heading('1. Executive Summary', level=1)
    
    summary_para = doc.add_paragraph()
    summary_para.add_run(f'This report provides a comprehensive analysis of {company_name} based on extracted financial data from annual reports. ')
    summary_para.add_run('The analysis includes historical EPS trends, ROE performance, market sentiment, and strategic insights.')
    
    doc.add_paragraph()

def add_market_sentiment_analysis(doc: Document, company_name: str, industry: str, research_data: Dict[str, Any], company_info: Dict[str, Any]):
    """Add market sentiment analysis section"""
    doc.add_heading('2. Market Sentiment Analysis', level=1)
    
    # Market sentiment - Prefer richer 'market_analysis' over default 'market_sentiment'
    market_sentiment_data = None
    if research_data and 'market_analysis' in research_data and research_data['market_analysis']:
        market_sentiment_data = research_data['market_analysis']
    elif research_data and 'market_sentiment' in research_data:
        market_sentiment_data = research_data['market_sentiment']
    
    if market_sentiment_data:
        # Overall sentiment
        if 'overall_sentiment' in market_sentiment_data:
            overall_para = doc.add_paragraph()
            overall_para.add_run('Overall Sentiment: ').bold = True
            overall_para.add_run(market_sentiment_data['overall_sentiment'])
        
        # P/E ratio analysis
        if 'pe_ratio_analysis' in market_sentiment_data:
            pe_para = doc.add_paragraph()
            pe_para.add_run('P/E Ratio Analysis: ').bold = True
            pe_para.add_run(market_sentiment_data['pe_ratio_analysis'])
        
        # Market position
        if 'market_position' in market_sentiment_data:
            position_para = doc.add_paragraph()
            position_para.add_run('Market Position: ').bold = True
            position_para.add_run(market_sentiment_data['market_position'])
        
        # Macro sentiment (fallback for old format)
        if 'macro_sentiment' in market_sentiment_data:
            macro_para = doc.add_paragraph()
            macro_para.add_run('Macro Sentiment: ').bold = True
            macro_para.add_run(market_sentiment_data['macro_sentiment'])
        
        # Micro sentiment (fallback for old format)
        if 'micro_sentiment' in market_sentiment_data:
            micro_para = doc.add_paragraph()
            micro_para.add_run('Micro Sentiment: ').bold = True
            micro_para.add_run(market_sentiment_data['micro_sentiment'])
        
        # Key insights (fallback for old format)
        if 'key_insights' in market_sentiment_data:
            insights_para = doc.add_paragraph()
            insights_para.add_run('Key Insights:').bold = True
            for insight in market_sentiment_data['key_insights']:
                doc.add_paragraph(insight, style='List Bullet')
    else:
        doc.add_paragraph('Market sentiment analysis not available.')
        
        # Show agent messages if available (but clean them first)
        if research_data and 'agent_messages' in research_data:
            doc.add_paragraph()
            doc.add_paragraph('Agent Messages:', style='Heading 3')
            for agent_name, message in research_data['agent_messages'].items():
                # Clean the message to remove thinking process
                cleaned_message = clean_agent_message(message)
                doc.add_paragraph(f'{agent_name}:', style='Heading 4')
                doc.add_paragraph(cleaned_message)
                doc.add_paragraph()
    
    doc.add_paragraph()

def add_swot_analysis(doc: Document, company_name: str, industry: str, research_data: Dict[str, Any], company_info: Dict[str, Any]):
    """Add SWOT analysis section"""
    doc.add_heading('3. SWOT Analysis', level=1)
    
    # Debug: Log what we received
    print(f"DEBUG SWOT: research_data keys: {list(research_data.keys()) if research_data else 'None'}")
    if research_data and 'swot_analysis' in research_data:
        print(f"DEBUG SWOT: swot_analysis keys: {list(research_data['swot_analysis'].keys())}")
        print(f"DEBUG SWOT: strengths count: {len(research_data['swot_analysis'].get('strengths', []))}")
    
    if research_data and 'swot_analysis' in research_data:
        swot = research_data['swot_analysis']
        
        # Check if SWOT analysis has meaningful content
        has_content = False
        content_count = 0
        for key in ['strengths', 'weaknesses', 'opportunities', 'threats']:
            if key in swot and swot[key] and len(swot[key]) > 0:
                has_content = True
                content_count += len(swot[key])
        
        print(f"DEBUG SWOT: has_content={has_content}, content_count={content_count}")
        
        if has_content and content_count >= 8:  # At least 2 items per category
            # Strengths
            if 'strengths' in swot and swot['strengths']:
                doc.add_heading('Strengths', level=2)
                for strength in swot['strengths']:
                    doc.add_paragraph(strength, style='List Bullet')
            else:
                doc.add_heading('Strengths', level=2)
                doc.add_paragraph('No specific strengths identified in the analysis.')
            
            # Weaknesses
            if 'weaknesses' in swot and swot['weaknesses']:
                doc.add_heading('Weaknesses', level=2)
                for weakness in swot['weaknesses']:
                    doc.add_paragraph(weakness, style='List Bullet')
            else:
                doc.add_heading('Weaknesses', level=2)
                doc.add_paragraph('No specific weaknesses identified in the analysis.')
            
            # Opportunities
            if 'opportunities' in swot and swot['opportunities']:
                doc.add_heading('Opportunities', level=2)
                for opportunity in swot['opportunities']:
                    doc.add_paragraph(opportunity, style='List Bullet')
            else:
                doc.add_heading('Opportunities', level=2)
                doc.add_paragraph('No specific opportunities identified in the analysis.')
            
            # Threats
            if 'threats' in swot and swot['threats']:
                doc.add_heading('Threats', level=2)
                for threat in swot['threats']:
                    doc.add_paragraph(threat, style='List Bullet')
            else:
                doc.add_heading('Threats', level=2)
                doc.add_paragraph('No specific threats identified in the analysis.')
        else:
            # Generate fallback SWOT analysis based on industry and company info
            doc.add_paragraph('SWOT analysis was attempted but insufficient meaningful content was generated.')
            doc.add_paragraph('Generating fallback analysis based on industry and company information...')
            
            # Generate industry-specific fallback SWOT
            fallback_swot = generate_fallback_swot_analysis(company_name, industry)
            
            if fallback_swot:
                # Strengths
                doc.add_heading('Strengths', level=2)
                for strength in fallback_swot.get('strengths', []):
                    doc.add_paragraph(strength, style='List Bullet')
                
                # Weaknesses
                doc.add_heading('Weaknesses', level=2)
                for weakness in fallback_swot.get('weaknesses', []):
                    doc.add_paragraph(weakness, style='List Bullet')
                
                # Opportunities
                doc.add_heading('Opportunities', level=2)
                for opportunity in fallback_swot.get('opportunities', []):
                    doc.add_paragraph(opportunity, style='List Bullet')
                
                # Threats
                doc.add_heading('Threats', level=2)
                for threat in fallback_swot.get('threats', []):
                    doc.add_paragraph(threat, style='List Bullet')
                
                doc.add_paragraph()
                doc.add_paragraph('Note: This SWOT analysis was generated as a fallback based on industry knowledge and company information. For more accurate analysis, please ensure comprehensive financial data is available.')
            else:
                doc.add_paragraph('Unable to generate fallback SWOT analysis. Please review the source data and try again.')
            # Intentionally omit raw agent messages from SWOT section
    else:
        doc.add_paragraph('SWOT analysis not available.')
        
        # Generate fallback SWOT analysis
        doc.add_paragraph('Generating fallback analysis based on industry and company information...')
        
        fallback_swot = generate_fallback_swot_analysis(company_name, industry)
        
        if fallback_swot:
            # Strengths
            doc.add_heading('Strengths', level=2)
            for strength in fallback_swot.get('strengths', []):
                doc.add_paragraph(strength, style='List Bullet')
            
            # Weaknesses
            doc.add_heading('Weaknesses', level=2)
            for weakness in fallback_swot.get('weaknesses', []):
                doc.add_paragraph(weakness, style='List Bullet')
            
            # Opportunities
            doc.add_heading('Opportunities', level=2)
            for opportunity in fallback_swot.get('opportunities', []):
                doc.add_paragraph(opportunity, style='List Bullet')
            
            # Threats
            doc.add_heading('Threats', level=2)
            for threat in fallback_swot.get('threats', []):
                doc.add_paragraph(threat, style='List Bullet')
            
            doc.add_paragraph()
            doc.add_paragraph('Note: This SWOT analysis was generated as a fallback based on industry knowledge and company information. For more accurate analysis, please ensure comprehensive financial data is available.')
        else:
            doc.add_paragraph('Unable to generate fallback SWOT analysis. Please review the source data and try again.')
        # Intentionally omit raw agent messages from SWOT section
    
    doc.add_paragraph()

def generate_fallback_swot_analysis(company_name: str, industry: str) -> Dict[str, List[str]]:
    """Generate fallback SWOT analysis based on industry and company information"""
    
    # Industry-specific SWOT templates
    industry_swot_templates = {
        'transportation': {
            'strengths': [
                f'Established network and infrastructure for {company_name}',
                'Specialization in perishable goods handling',
                'Robust technology integration and operational efficiency',
                'Strong market position in logistics and supply chain'
            ],
            'weaknesses': [
                'High fuel costs and operational expenses',
                'Regulatory complexities in cross-border operations',
                'Dependency on economic cycles and trade volumes',
                'Market sensitivity to global economic conditions'
            ],
            'opportunities': [
                'Technology-driven services and automation',
                'Expansion into new markets and routes',
                'Strategic partnerships with key industry players',
                'Sustainability initiatives and green logistics'
            ],
            'threats': [
                'Rising fuel costs impacting operational expenses',
                'Regulatory changes and compliance requirements',
                'Competition from tech disruptors and new entrants',
                'Economic downturns affecting trade volumes'
            ]
        },
        'banking': {
            'strengths': [
                f'Established customer base and brand recognition for {company_name}',
                'Diversified revenue streams and financial stability',
                'Strong regulatory compliance framework',
                'Extensive branch network and digital presence'
            ],
            'weaknesses': [
                'High regulatory compliance costs',
                'Dependency on interest rate environment',
                'Cybersecurity risks and operational complexity',
                'Market sensitivity to economic conditions'
            ],
            'opportunities': [
                'Digital transformation and fintech partnerships',
                'Expansion into new markets and customer segments',
                'Cross-selling opportunities to existing customers',
                'Cost optimization through automation and AI'
            ],
            'threats': [
                'Increasing competition from fintech companies',
                'Regulatory changes and compliance requirements',
                'Economic downturns and credit risks',
                'Cybersecurity threats and data breaches'
            ]
        },
        'technology': {
            'strengths': [
                f'Innovation and R&D capabilities for {company_name}',
                'Scalable business model with global reach',
                'Strong intellectual property and market position',
                'Technology leadership and expertise'
            ],
            'weaknesses': [
                'High R&D costs impacting profit margins',
                'Rapid technology obsolescence risk',
                'Dependency on key personnel and expertise',
                'Market sensitivity to technological changes'
            ],
            'opportunities': [
                'Emerging technologies (AI, IoT, Cloud) market expansion',
                'International market expansion opportunities',
                'Strategic partnerships and acquisitions',
                'Subscription-based revenue models for stability'
            ],
            'threats': [
                'Intense competition and market disruption',
                'Rapid technological changes and obsolescence',
                'Cybersecurity threats and data privacy concerns',
                'Regulatory scrutiny and compliance costs'
            ]
        }
    }
    
    # Get industry-specific template or use general template
    industry_lower = industry.lower()
    swot_template = None
    
    for key, template in industry_swot_templates.items():
        if key in industry_lower:
            swot_template = template
            break
    
    if not swot_template:
        # General template for unknown industries
        swot_template = {
            'strengths': [
                f'Established market position for {company_name}',
                'Strong brand recognition and customer base',
                'Operational efficiency and market presence',
                'Industry expertise and experience'
            ],
            'weaknesses': [
                'Market competition and operational costs',
                'Dependency on external factors and market conditions',
                'Resource constraints and scalability challenges',
                'Market sensitivity to economic conditions'
            ],
            'opportunities': [
                'Market expansion opportunities and growth potential',
                'Digital transformation and technology adoption',
                'Strategic partnerships and collaborations',
                'Product/service innovation and diversification'
            ],
            'threats': [
                'Economic uncertainty and market volatility',
                'Regulatory changes and compliance costs',
                'Competitive pressures and market disruption',
                'Technology disruption and changing customer preferences'
            ]
        }
    
    return swot_template

def add_financial_analysis(doc: Document, eps_data: Dict[str, Any], roe_data: Dict[str, Any], calculations: Dict[str, Any], company_name: str, company_info: Dict[str, Any]):
    """
    Adds the Financial Performance and Calculations section to the document.
    """
    doc.add_heading('4. Financial Performance', level=1)
    
    # EPS Data Table (REMOVED PE COLUMN)
    if eps_data:
        doc.add_heading('Earnings Per Share (EPS) Data', level=2)
        
        # Determine the most recent up to 10 years
        eps_year_keys = sorted([int(y) for y in eps_data.keys() if str(y).isdigit()])
        eps_recent_years = [str(y) for y in eps_year_keys[-10:]] if eps_year_keys else []
        years_with_data = len([y for y in eps_recent_years if isinstance(eps_data.get(y, {}), dict) and eps_data.get(y, {}).get('basic_eps') is not None])
        total_years = len(eps_recent_years)
        
        # Add note about 10-year filtering
        note_para = doc.add_paragraph()
        note_para.add_run('Note: ').bold = True
        note_para.add_run(f'This analysis shows the most recent {total_years} years of EPS data (up to 10 years) extracted from the annual report. All available historical data was extracted and analyzed.')
        note_para.style.font.size = Pt(10)
        note_para.style.font.italic = True
        doc.add_paragraph()
        
        # Add units information
        if 'units' in eps_data:
            units_para = doc.add_paragraph()
            eps_unit = eps_data['units'].get('eps_unit', 'Unknown')
            currency = eps_data['units'].get('currency', 'Unknown')
            units_para.add_run(f'Units: ').bold = True
            units_para.add_run(f'EPS in {eps_unit}, Currency: {currency}')
            doc.add_paragraph()
        
        # Create table (REMOVED PE COLUMN)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Year'
        # Use currency from units data
        currency = eps_data['units'].get('currency', 'SGD') if 'units' in eps_data else 'SGD'
        header_cells[1].text = f'Basic EPS ({currency})'
        
        # Track corrections for footnotes
        corrections = []
        correction_counter = 1
        
        # Add data rows (most recent up to 10 years)
        for year in eps_recent_years:
            # Skip non-year keys like 'units'
            if not year.isdigit():
                continue
                
            row_cells = table.add_row().cells
            row_cells[0].text = str(year)
            
            year_data = eps_data[year]
            if isinstance(year_data, dict):
                basic_eps = year_data.get('basic_eps')
                
                # Convert cents to dollars if needed
                eps_unit = eps_data.get('units', {}).get('eps_unit', 'cents')
                if eps_unit == 'cents' and basic_eps is not None:
                    # Convert cents to dollars
                    basic_eps = basic_eps / 100
                
                # Check if this value was corrected
                if 'original_extracted' in year_data and 'replaced_reason' in year_data:
                    # This value was corrected - show corrected value with footnote
                    corrected_value = f"{basic_eps:.2f}" if basic_eps is not None else 'N/A'
                    row_cells[1].text = f"{corrected_value}¹"
                    
                    # Store correction info for footnote
                    original_value = year_data.get('original_extracted')
                    reason = year_data.get('replaced_reason')
                    corrections.append(f"¹ {year}: Corrected from {original_value} to {corrected_value} ({reason})")
                else:
                    # This value was not corrected
                    row_cells[1].text = f"{basic_eps:.2f}" if basic_eps is not None else 'N/A'
            else:
                # Convert cents to dollars if needed
                eps_unit = eps_data.get('units', {}).get('eps_unit', 'cents')
                if eps_unit == 'cents' and year_data is not None:
                    year_data = year_data / 100
                row_cells[1].text = f"{year_data:.2f}" if year_data else 'N/A'
        
        # Add footnotes for corrections
        if corrections:
            doc.add_paragraph()
            footnote_heading = doc.add_heading('Data Corrections', level=3)
            footnote_heading.style.font.size = Pt(10)
            
            for correction in corrections:
                footnote_para = doc.add_paragraph(correction, style='Normal')
                footnote_para.style.font.size = Pt(9)
                footnote_para.style.font.italic = True
    
    # Add EPS Graph
    if eps_data:
        doc.add_heading('EPS Historical Trend Chart', level=2)
        
        # Create the graph - filter to most recent up to 10 years
        eps_data_filtered = {k: eps_data[k] for k in eps_recent_years if k.isdigit()}
        graph_buffer = create_eps_graph(eps_data_filtered, eps_data.get('units', {}))
        
        if graph_buffer:
            # Add the graph to the document
            doc.add_picture(graph_buffer, width=Inches(6))
            
            # Add caption
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            years_with_data = len([y for y, d in eps_data_filtered.items() if isinstance(d, dict) and d.get('basic_eps') is not None])
            caption.add_run(f'Figure 1: {company_name} - {total_years}-Year Historical Basic Earnings Per Share Trend').italic = True
            
            # Add spacing
            doc.add_paragraph()
        else:
            doc.add_paragraph('Graph could not be generated due to insufficient data.')
    
    # Add ROE Data and Graph
    if roe_data:
        doc.add_heading('Return on Equity (ROE) Data', level=2)
        
        # Determine the most recent up to 10 years
        roe_year_keys = sorted([int(y) for y in roe_data.keys() if str(y).isdigit()])
        roe_recent_years = [str(y) for y in roe_year_keys[-10:]] if roe_year_keys else []
        roe_years_with_data = len([y for y in roe_recent_years if isinstance(roe_data.get(y, {}), dict) and roe_data.get(y, {}).get('basic_roe') is not None])
        roe_note_para = doc.add_paragraph()
        roe_note_para.add_run('Note: ').bold = True
        roe_note_para.add_run(f'This analysis shows the most recent {len(roe_recent_years)} years of ROE data (up to 10 years) extracted from the annual report. All available historical data was extracted and analyzed.')
        roe_note_para.style.font.size = Pt(10)
        roe_note_para.style.font.italic = True
        doc.add_paragraph()
        
        # Add units information for ROE
        if 'units' in roe_data:
            roe_units_para = doc.add_paragraph()
            roe_unit = roe_data['units'].get('roe_unit', 'Unknown')
            roe_units_para.add_run(f'ROE Units: ').bold = True
            roe_units_para.add_run(f'{roe_unit}')
            doc.add_paragraph()
        
        # Create ROE table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Year'
        header_cells[1].text = 'ROE (%)'
        
        # Add data rows (most recent up to 10 years)
        for year in roe_recent_years:
            # Skip non-year keys like 'units'
            if not year.isdigit():
                continue
                
            row_cells = table.add_row().cells
            row_cells[0].text = str(year)
            
            year_data = roe_data[year]
            if isinstance(year_data, dict):
                # Handle structured ROE data
                basic_roe = year_data.get('basic_roe')
                if basic_roe is not None:
                    row_cells[1].text = f"{basic_roe:.2f}%"
                else:
                    row_cells[1].text = 'N/A'
            else:
                # Handle direct ROE value
                if year_data is not None:
                    row_cells[1].text = f"{year_data:.2f}%"
                else:
                    row_cells[1].text = 'N/A'
        
        doc.add_paragraph()
        
        # Add ROE Graph
        doc.add_heading('ROE Historical Trend Chart', level=2)
        
        # Create the ROE graph - filter to most recent up to 10 years
        roe_data_filtered = {k: roe_data[k] for k in roe_recent_years if k.isdigit()}
        roe_graph_buffer = create_roe_graph(roe_data_filtered, roe_data.get('units', {}))
        
        if roe_graph_buffer:
            # Add the graph to the document
            doc.add_picture(roe_graph_buffer, width=Inches(6))
            
            # Add caption
            roe_caption = doc.add_paragraph()
            roe_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            roe_caption.add_run(f'Figure 2: {company_name} - {len(roe_recent_years)}-Year Historical Return on Equity Trend').italic = True
            
            # Add spacing
            doc.add_paragraph()
        else:
            doc.add_paragraph('ROE graph could not be generated due to insufficient data.')
    
    # Financial Calculations & Valuation section removed per user request

def add_validation_and_valuation(doc: Document, validation_results: Dict[str, Any], valuation_data: Dict[str, Any], calculations: Dict[str, Any], market_data: Dict[str, Any], company_name: str, company_info: Dict[str, Any], eps_data: Dict[str, Any] = None):
    """Merged: Data Validation Results & Recommendations + Valuation Analysis"""
    doc.add_heading('5. Data Validation Results, Recommendations & Valuation', level=1)
    
    if not validation_results:
        doc.add_paragraph('No validation results available for this analysis.')
        doc.add_paragraph()
        # Do not return here; still show the Valuation Analysis subsection below
    
    # Overall confidence score
    overall_confidence = validation_results.get('overall_confidence', 0.0)
    doc.add_heading('Overall Validation Confidence', level=2)
    
    confidence_para = doc.add_paragraph()
    confidence_para.add_run('Confidence Score: ').bold = True
    confidence_percentage = overall_confidence * 100
    
    if confidence_percentage >= 80:
        confidence_para.add_run(f'{confidence_percentage:.1f}% (High Confidence)').font.color.rgb = RGBColor(0, 128, 0)  # Green
    elif confidence_percentage >= 60:
        confidence_para.add_run(f'{confidence_percentage:.1f}% (Moderate Confidence)').font.color.rgb = RGBColor(255, 165, 0)  # Orange
    else:
        confidence_para.add_run(f'{confidence_percentage:.1f}% (Low Confidence)').font.color.rgb = RGBColor(255, 0, 0)  # Red
    
    # Validation sources
    validation_sources = validation_results.get('validation_sources', {})
    if validation_sources:
        doc.add_heading('Validation Sources', level=2)
        sources_para = doc.add_paragraph()
        sources_para.add_run('External validation was performed using the following sources: ').bold = True
        
        source_names = []
        for source_name, source_data in validation_sources.items():
            if source_data and 'confidence' in source_data:
                confidence = source_data['confidence'] * 100
                source_names.append(f'{source_name.replace("_", " ").title()} ({confidence:.1f}% confidence)')
        
        sources_para.add_run(', '.join(source_names))
    
    # Discrepancies
    discrepancies = validation_results.get('discrepancies', [])
    if discrepancies:
        doc.add_heading('Data Discrepancies Found', level=2)
        doc.add_paragraph('The following discrepancies were identified between extracted data and external sources:')
        
        # Create table for discrepancies
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Year'
        header_cells[1].text = 'Source'
        header_cells[2].text = 'Extracted Value'
        header_cells[3].text = 'External Value'
        header_cells[4].text = 'Difference (%)'
        
        # Add discrepancy rows
        for discrepancy in discrepancies:
            row_cells = table.add_row().cells
            row_cells[0].text = str(discrepancy.get('year', 'N/A'))
            row_cells[1].text = discrepancy.get('source', 'N/A').replace('_', ' ').title()
            
            # Get extracted and external values
            discrepancy_data = discrepancy.get('discrepancy', {})
            extracted_value = discrepancy_data.get('extracted', 'N/A')
            external_value = discrepancy_data.get('yahoo', discrepancy_data.get('alpha_vantage', discrepancy_data.get('polygon', 'N/A')))
            diff_percent = discrepancy_data.get('difference_percent', 'N/A')
            
            row_cells[2].text = f"{extracted_value:.2f}" if isinstance(extracted_value, (int, float)) else str(extracted_value)
            row_cells[3].text = f"{external_value:.2f}" if isinstance(external_value, (int, float)) else str(external_value)
            row_cells[4].text = f"{diff_percent:.1f}%" if isinstance(diff_percent, (int, float)) else str(diff_percent)
    
    # Recommendations
    recommendations = validation_results.get('recommendations', [])
    if recommendations:
        doc.add_heading('Validation Recommendations', level=2)
        doc.add_paragraph('Based on the validation analysis, the following recommendations are provided:')
        
        for recommendation in recommendations:
            rec_para = doc.add_paragraph()
            rec_para.add_run('• ').bold = True
            rec_para.add_run(recommendation)
    
    # Action plan
    action_plan = validation_results.get('action_plan', '')
    if action_plan:
        doc.add_heading('Action Plan', level=2)
        action_para = doc.add_paragraph()
        action_para.add_run('Recommended Action: ').bold = True
        
        if action_plan == 'use_extracted_data':
            action_para.add_run('Use extracted data as primary source (high confidence)')
        elif action_plan == 'use_extracted_data_with_warnings':
            action_para.add_run('Use extracted data with warnings for discrepancies')
        elif action_plan == 'use_external_sources':
            action_para.add_run('Use external sources for years with high discrepancies')
        elif action_plan == 'recheck_annual_report':
            action_para.add_run('Re-check annual report for accuracy')
        elif action_plan == 'manual_review':
            action_para.add_run('Manual review required - data inconsistencies detected')
        else:
            action_para.add_run(action_plan)
    
    # Key Calculations moved into the Valuation Analysis table below

    # Data quality summary
    doc.add_heading('Data Quality Summary', level=2)
    quality_para = doc.add_paragraph()
    
    if overall_confidence >= 0.8:
        quality_para.add_run('OK: ').bold = True
        quality_para.add_run('High quality data - extracted values match external sources well')
    elif overall_confidence >= 0.6:
        quality_para.add_run('WARN: ').bold = True
        quality_para.add_run('Moderate quality data - some discrepancies found, review recommended')
    else:
        quality_para.add_run('LOW: ').bold = True
        quality_para.add_run('Low quality data - significant discrepancies found, manual review required')
    
    # Valuation Analysis (merged + includes calculations)
    doc.add_heading('Valuation Analysis', level=2)
    has_calc = calculations and isinstance(calculations, Dict) and len(calculations) > 0
    has_val = valuation_data and isinstance(valuation_data, Dict) and any(v is not None for v in valuation_data.values())
    if has_calc or has_val:
        val_table = doc.add_table(rows=1, cols=2)
        val_table.style = 'Table Grid'
        header_cells = val_table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'

        def _fmt_number(val: Any) -> str:
            if val is None:
                return 'N/A'
            return f"{float(val):.2f}" if isinstance(val, (int, float)) else str(val)

        # Get currency from eps_data
        currency = 'SGD'  # default
        if eps_data and 'units' in eps_data:
            currency = eps_data['units'].get('currency', 'SGD')
        
        def _add_row(name: str, value: Any, kind: str = 'number'):
            row_cells = val_table.add_row().cells
            row_cells[0].text = name
            if value is None:
                row_cells[1].text = 'N/A'
                return
            if kind == 'percent':
                row_cells[1].text = f"{float(value):.2f}%"
            elif kind == 'currency':
                row_cells[1].text = f"{float(value):.2f} {currency}"
            else:
                row_cells[1].text = _fmt_number(value)

        # Add calculation metrics (single Current EPS occurrence)
        if has_calc:
            # Only show values when numerically valid
            cagr_val = calculations.get('cagr') if isinstance(calculations.get('cagr'), (int, float)) else None
            avg_pe_val = calculations.get('avg_pe_ratio') if isinstance(calculations.get('avg_pe_ratio'), (int, float)) else None
            low_pe_val = calculations.get('lowest_pe_ratio') if isinstance(calculations.get('lowest_pe_ratio'), (int, float)) else None
            curr_eps_val = calculations.get('current_eps') if isinstance(calculations.get('current_eps'), (int, float)) else None
            proj_eps_val = calculations.get('projected_eps_10yr') if isinstance(calculations.get('projected_eps_10yr'), (int, float)) else None

            _add_row('CAGR', cagr_val, 'percent')
            _add_row('Average PE Ratio', avg_pe_val)
            _add_row('Lowest PE Ratio', low_pe_val)
            _add_row('Current EPS', curr_eps_val, 'currency')
            _add_row('Projected EPS 10Yr', proj_eps_val, 'currency')

        # Add valuation metrics, excluding duplicates already shown
        ordered_keys = [
            # First: keep calculations already added above
            # Then future prices (nominal): avg PE then lowest PE
            'nominal_future_price', 'nominal_future_price_lowest',
            # Source and WACC
            'wacc_source', 'discount_rate_percent',
            # PVs
            'pv_base', 'pv_lowest_pe',
            # Safety margins
            'safety_margin_low_percent', 'safety_margin_high_percent',
            # Final PV after safety margins
            'nominal_pv_5', 'nominal_pv_20'
        ]
        if has_val:
            for key in ordered_keys:
                if key not in valuation_data:
                    continue
                val = valuation_data.get(key)
                # Guard: Only show numeric valuation numbers where appropriate
                if key in ['nominal_future_price', 'nominal_future_price_lowest', 'pv_base', 'pv_lowest_pe', 'nominal_pv_5', 'nominal_pv_20']:
                    if not isinstance(val, (int, float)):
                        val = None
                lk = key.lower()
                # Custom rename map for display labels
                if key == 'pv_base':
                    display = 'Future Price(PV - Avg PE)'
                elif key == 'pv_lowest_pe':
                    display = 'Future Price(PV - Lowest PE)'
                elif key == 'nominal_pv_5':
                    display = 'Estimated Intrinsic Value (Best Case)'
                elif key == 'nominal_pv_20':
                    display = 'Estimated Intrinsic Value (Worst Case)'
                elif key == 'nominal_future_price':
                    display = 'Future Price (10 Years - Avg PE ratio)'
                elif key == 'nominal_future_price_lowest':
                    display = 'Future Price (10 Years - Lowest PE ratio)'
                else:
                    base = key.replace('_', ' ').title()
                    # Apply specific capitalization rules without affecting 'Percent'
                    if ('pe' in lk) and (not lk.endswith('percent')):
                        display = base.replace('Pe', 'PE')
                    else:
                        display = base
                    display = display.replace('Pv', 'PV').replace('Eps', 'EPS')
                if key == 'wacc_source':
                    _add_row('WACC Source', str(val).title() if val is not None else None)
                elif lk.endswith('percent'):
                    display = display.replace(' Percent', ' (%)')
                    _add_row(display, val, 'percent')
                elif any(tag in lk for tag in ['eps', 'price', 'pv']):
                    _add_row(display, val, 'currency')
                else:
                    _add_row(display, val)
        doc.add_paragraph()
        doc.add_heading('Valuation Insights', level=3)
        for insight in generate_valuation_insights(valuation_data or {}, company_name, eps_data):
            doc.add_paragraph(insight, style='List Bullet')
    else:
        doc.add_paragraph('Valuation analysis not available.')
        doc.add_paragraph('Provide inputs or enable data sources to compute valuation metrics.')

def add_valuation_analysis(doc: Document, valuation_data: Dict[str, Any], market_data: Dict[str, Any], company_name: str, company_info: Dict[str, Any]):
    """Add valuation analysis section"""
    doc.add_heading('6. Valuation Analysis', level=1)
    
    if valuation_data and len(valuation_data) > 0:
        # Check if we have meaningful valuation data
        meaningful_data = False
        for key, value in valuation_data.items():
            if value is not None and value != 0:
                meaningful_data = True
                break
        
        if meaningful_data:
            # Create valuation table
            val_table = doc.add_table(rows=1, cols=2)
            val_table.style = 'Table Grid'
            
            # Header row
            header_cells = val_table.rows[0].cells
            header_cells[0].text = 'Valuation Metric'
            header_cells[1].text = 'Value (SGD)'
            
            # Add valuation data
            for key, value in valuation_data.items():
                row_cells = val_table.add_row().cells
                
                # Format the metric name with proper capitalization
                metric_name = key.replace("_", " ").title()
                # Apply specific capitalization rules for valuation metrics
                if 'pe' in key.lower():
                    metric_name = metric_name.replace('Pe', 'PE')
                if 'eps' in key.lower():
                    metric_name = metric_name.replace('Eps', 'EPS')
                if 'pv' in key.lower():
                    metric_name = metric_name.replace('Pv', 'PV')
                row_cells[0].text = metric_name
                
                # Format the value
                if value is not None and value != 0:
                    if isinstance(value, (int, float)):
                        row_cells[1].text = f"{value:.2f}"
                    else:
                        row_cells[1].text = str(value)
                else:
                    row_cells[1].text = "N/A"
            
            doc.add_paragraph()
            
            # Add valuation insights
            doc.add_heading('Valuation Insights', level=2)
            
            # Analyze the valuation data and provide insights
            insights = generate_valuation_insights(valuation_data, company_name, None)
            for insight in insights:
                doc.add_paragraph(f'• {insight}', style='List Bullet')
        else:
            doc.add_paragraph('Valuation analysis data is incomplete or unavailable.')
            doc.add_paragraph('This may be due to insufficient financial data or calculation errors.')
            doc.add_paragraph('Recommendations:')
            doc.add_paragraph('• Ensure comprehensive EPS and PE ratio data is available', style='List Bullet')
            doc.add_paragraph('• Verify that market data sources are accessible', style='List Bullet')
            doc.add_paragraph('• Check that financial calculations are properly configured', style='List Bullet')
    else:
        doc.add_paragraph('Valuation analysis not available.')
        doc.add_paragraph('This section requires comprehensive financial data including:')
        doc.add_paragraph('• Historical EPS data', style='List Bullet')
        doc.add_paragraph('• PE ratio information', style='List Bullet')
        doc.add_paragraph('• Market price data', style='List Bullet')
        doc.add_paragraph('• Growth rate calculations', style='List Bullet')
        
        doc.add_paragraph('To generate valuation analysis:')
        doc.add_paragraph('• Ensure PDF reports contain sufficient financial data', style='List Bullet')
        doc.add_paragraph('• Verify market data integration is working properly', style='List Bullet')
        doc.add_paragraph('• Check that all required calculations can be performed', style='List Bullet')
    
    doc.add_paragraph()

def generate_valuation_insights(valuation_data: Dict[str, Any], company_name: str, eps_data: Dict[str, Any] = None) -> List[str]:
    """Generate insights based on valuation data"""
    insights = []
    
    # Get currency from eps_data if available
    currency = 'USD'  # default
    if eps_data and 'units' in eps_data:
        currency = eps_data['units'].get('currency', 'USD')
    
    try:
        # Check for key valuation metrics
        if 'nominal_future_price' in valuation_data and valuation_data['nominal_future_price']:
            future_price = valuation_data['nominal_future_price']
            insights.append(f"Projected future price: {future_price:.2f} {currency} based on growth assumptions")
        
        if 'final_nominal_price' in valuation_data and valuation_data['final_nominal_price']:
            final_price = valuation_data['final_nominal_price']
            insights.append(f"Final valuation estimate: {final_price:.2f} {currency} considering multiple scenarios")
        
        if 'nominal_pv_5' in valuation_data and valuation_data['nominal_pv_5']:
            pv_5 = valuation_data['nominal_pv_5']
            insights.append(f"Estimated intrinsic value (best case scenario): {pv_5:.2f} {currency}")
        
        if 'nominal_pv_20' in valuation_data and valuation_data['nominal_pv_20']:
            pv_20 = valuation_data['nominal_pv_20']
            insights.append(f"Estimated intrinsic value (worst case scenario): {pv_20:.2f} {currency}")
        
        # Add general insights if no specific data
        if not insights:
            insights.append("Valuation analysis requires comprehensive financial data and market information")
            insights.append("Consider supplementing with additional market data sources")
            insights.append("Review financial calculations and assumptions for accuracy")
    
    except Exception as e:
        insights.append(f"Error generating valuation insights: {str(e)}")
        insights.append("Please review the valuation data and calculations")
    
    return insights

def create_valuation_graph(market_prices: List[float], company_name: str):
    """Create valuation graph"""
    try:
        if len(market_prices) < 2:
            return None
        
        years = list(range(len(market_prices)))
        
        plt.figure(figsize=(10, 6))
        plt.plot(years, market_prices, marker='o', linewidth=2, markersize=6)
        
        plt.title('Projected Stock Price Trend', fontsize=14, fontweight='bold')
        plt.xlabel('Years', fontsize=12)
        plt.ylabel('Price (SGD)', fontsize=12)
        plt.grid(True, alpha=0.3)
        
        plt.tight_layout()
        
        buffer = BytesIO()
        plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
        buffer.seek(0)
        plt.close()
        
        return buffer
        
    except Exception as e:
        print(f"Error creating valuation graph: {e}")
        return None

def add_market_analysis(doc: Document, company_name: str, ticker: str, industry: str, research_data: Dict[str, Any], company_info: Dict[str, Any]):
    """Add market analysis section"""
    doc.add_heading('6. Market Analysis', level=1)
    
    if research_data:
        # Industry trends
        if 'industry_trends' in research_data:
            trends = research_data['industry_trends']
            
            if 'key_trends' in trends:
                doc.add_heading('Key Industry Trends', level=2)
                for trend in trends['key_trends']:
                    doc.add_paragraph(str(trend), style='List Bullet')
            
            if 'growth_drivers' in trends:
                doc.add_heading('Growth Drivers', level=2)
                for driver in trends['growth_drivers']:
                    doc.add_paragraph(str(driver), style='List Bullet')
            
            if 'challenges' in trends:
                doc.add_heading('Industry Challenges', level=2)
                for challenge in trends['challenges']:
                    doc.add_paragraph(str(challenge), style='List Bullet')
        
        # Competitive analysis
        if 'competitive_analysis' in research_data:
            comp_analysis = research_data['competitive_analysis']
            
            if 'main_competitors' in comp_analysis:
                doc.add_heading('Main Competitors', level=2)
                for competitor in comp_analysis['main_competitors']:
                    doc.add_paragraph(str(competitor), style='List Bullet')
            
            if 'competitive_advantages' in comp_analysis:
                doc.add_heading('Competitive Advantages', level=2)
                for advantage in comp_analysis['competitive_advantages']:
                    doc.add_paragraph(str(advantage), style='List Bullet')
    else:
        doc.add_paragraph('Market analysis not available.')
    
    doc.add_paragraph()

def add_conclusion(doc: Document, company_name: str, industry: str, research_data: Dict[str, Any], company_info: Dict[str, Any]):
    """Add conclusion section"""
    doc.add_heading('7. Conclusion', level=1)
    
    conclusion_para = doc.add_paragraph()
    conclusion_para.add_run(f'This analysis provides a comprehensive overview of {company_name} based on extracted financial data. ')
    conclusion_para.add_run('The historical performance, market position, and future projections have been evaluated to provide strategic insights.')
    
    doc.add_paragraph()

def add_sources(doc: Document, company_name: str, industry: str, research_data: Dict[str, Any], company_info: Dict[str, Any], eps_data: Dict[str, Any]):
    """Add sources section"""
    doc.add_heading('8. Sources', level=1)
    
    # Data sources
    doc.add_heading('Data Sources', level=2)
    doc.add_paragraph('• Annual Reports (PDF files)')
    doc.add_paragraph('• Company Financial Statements')
    doc.add_paragraph('• Market Research Data')
    
    # Analysis methodology
    doc.add_heading('Analysis Methodology', level=2)
    doc.add_paragraph('• AI-powered PDF extraction and analysis')
    doc.add_paragraph('• Historical trend analysis')
    doc.add_paragraph('• Industry benchmarking')
    doc.add_paragraph('• Financial ratio calculations')
    
    # Units and currency information
    if eps_data and 'units' in eps_data:
        doc.add_heading('Units and Currency', level=2)
        eps_unit = eps_data['units'].get('eps_unit', 'Unknown')
        currency = eps_data['units'].get('currency', 'Unknown')
        doc.add_paragraph(f'• EPS Unit: {eps_unit}')
        doc.add_paragraph(f'• Currency: {currency}')
    
    doc.add_paragraph()

if __name__ == "__main__":
    # Test the document creator
    test_eps_data = {
        "2024": {"basic_eps": 107, "pe_ratio": 12.0},
        "2023": {"basic_eps": 80, "pe_ratio": None},
        "2022": {"basic_eps": 112, "pe_ratio": None},
        "units": {"eps_unit": "cents", "currency": "SGD"}
    }
    
    test_roe_data = {
        "2024": {"basic_roe": 9.6},
        "2023": {"basic_roe": 7.6},
        "2022": {"basic_roe": 11.2},
        "units": {"roe_unit": "percentage"}
    }
    
    test_calculations = {
        "cagr": 6.25,
        "avg_pe_ratio": 18.0,
        "lowest_pe_ratio": 18.0,
        "current_eps": 107,
        "projected_eps_10yr": 201.3
    }
    
    create_word_document(
        eps_data=test_eps_data,
        roe_data=test_roe_data,
        market_data={},
        calculations=test_calculations,
        valuation_data={},
        research_data={}
    ) 